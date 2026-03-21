import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), str(m[1]))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, color=RGBColor(0x1B, 0x2A, 0x4A)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def create_concept_report():
    doc = Document()
    
    # 1. ТРАНСФОРМАЦИЯ (ДРАМАТИЧЕСКАЯ АРКА)
    add_styled_heading(doc, '1. ТРАНСФОРМАЦИЯ: ОТ ХАОСА К NEXUS PRIME', level=1)
    doc.add_paragraph(
        "Этот раздел описывает качественный скачок в подходе к разработке высокопроизводительных систем."
    )
    
    table_arc = doc.add_table(rows=1, cols=3)
    table_arc.style = 'Table Grid'
    hdr = table_arc.rows[0].cells
    hdr[0].text = 'Параметр'; hdr[1].text = 'БЫЛО (Ручной труд)'; hdr[2].text = 'СТАЛО (Nexus Protocol)'
    
    arc_data = [
        ['Разработка', 'Медленный бойлерплейт и ошибки типов', 'Агенты пишут Rust-ядро с проверкой памяти'],
        ['Производительность', 'Непредсказуемые задержки (GC)', 'Стабильные 100k msg/s (Zero-GC)'],
        ['Контекст', 'Рассеянные ADR и забытые решения', 'Глубокая компиляция истории в каждый файл']
    ]
    for p, b, a in arc_data:
        row = table_arc.add_row().cells
        row[0].text = p; row[1].text = b; row[2].text = a

    doc.add_paragraph("\n")

    # 2. ТЕМА И ЦЕЛЬ (Messaging Engine)
    add_styled_heading(doc, '2. ЦЕНТРАЛЬНАЯ ИДЕЯ: HIGH-PERF ENGINE', level=1)
    doc.add_paragraph(
        "Цель: Обработка 100,000 сообщений в секунду с задержкой <1ms на стандартном Windows-оборудовании (6GB RAM)."
    )
    
    # 3. ТЕХНИЧЕСКАЯ ФОРМУЛА
    add_styled_heading(doc, '3. ФОРМУЛА УСПЕХА (STEPS)', level=1)
    formula_box = doc.add_paragraph()
    formula_box.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = formula_box.add_run("RUST core + TOKIO async + NAMED PIPES (Windows IPC) = ZERO LATENCY SPIKES")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # 4. ВЫБРАННЫЙ СТЕК
    add_styled_heading(doc, '4. АРХИТЕКТУРНЫЙ СТЕК', level=1)
    stack_table = doc.add_table(rows=1, cols=2)
    stack_table.style = 'Table Grid'
    hdr = stack_table.rows[0].cells
    hdr[0].text = 'Компонент'; hdr[1].text = 'Решение'
    
    stack_data = [
        ['Движок (Core)', 'Rust 1.75+ (Tokio, Sled)'],
        ['API / Logic', 'Python 3.13 (FastHTML)'],
        ['Оркестрация', 'Go (PocketBase)'],
        ['Оптимизация', 'Nexus System Control (Turbo Boost)']
    ]
    for c, r in stack_data:
        row = stack_table.add_row().cells
        row[0].text = c; row[1].text = r

    # 5. ПЛАН РЕАЛИЗАЦИИ
    add_styled_heading(doc, '5. ДОРОЖНАЯ КАРТА (ROADMAP)', level=1)
    steps = [
        "Phase 1: Прототипирование Named Pipes IPC на Rust.",
        "Phase 2: Интеграция Python-контроллера для управления движком.",
        "Phase 3: Стресс-тестирование (100k/s) под нагрузкой Turbo Boost.",
        "Phase 4: Финализация документации и деплой."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # Save
    output_path = r'e:\Downloads\--ANTIGRAVITY store\IDE-optimus\PROJECT\outputs\Nexus_Concept_Report.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"CONCEPT_GENERATED: {output_path}")

if __name__ == "__main__":
    create_concept_report()
