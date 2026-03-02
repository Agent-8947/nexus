import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# DESIGN TOKENS
COLOR_DARK = RGBColor(0x0F, 0x17, 0x2A)   # Midnight
COLOR_PRIMARY = RGBColor(0x25, 0x63, 0xEB) # Royal Blue
COLOR_ACCENT = RGBColor(0x38, 0xBD, 0xF8) # Sky
PATH_LOGO = r'C:\Users\MAC\.gemini\antigravity\brain\858d0bd1-44f8-418a-847c-07d05808974d\nexus_premium_header_1772196955533.png'

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, start=150, bottom=100, end=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), str(m[1]))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_presentation():
    doc = Document()
    
    # 1. BRANDING HEADER
    if os.path.exists(PATH_LOGO):
        doc.add_picture(PATH_LOGO, width=Inches(6.0))
    
    title = doc.add_heading('NEXUS ECOSYSTEM: THE FUTURE OF AI OPS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Segoe UI Semibold'
        run.font.color.rgb = COLOR_DARK

    doc.add_paragraph("\n")

    # 2. TRANSFORMATIONAL ARC
    h1 = doc.add_heading('1. THE ARCHITECTURAL LEAP', level=1)
    for run in h1.runs: 
        run.font.name = 'Segoe UI'
        run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph('Moving beyond traditional bottlenecks into asynchronous agentic workflows.')
    
    table_arc = doc.add_table(rows=1, cols=3)
    table_arc.style = 'Table Grid'
    hdr = table_arc.rows[0].cells
    cols = ['State', 'Legacy Build', 'Nexus Prime']
    for i, c in enumerate(cols):
        hdr[i].text = c
        set_cell_background(hdr[i], '0F172A')
        run = hdr[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

    arc_data = [
        ['Dev Cycle', 'Manual overhead', 'Self-Healing AI Agents'],
        ['Scaling', 'Vertical limits', 'Bayesian High-Perf Mesh'],
        ['Control', 'Fractured CLI', 'Nexus Omni-Dashboard']
    ]
    for s, b, a in arc_data:
        row = table_arc.add_row().cells
        row[0].text = s; row[1].text = b; row[2].text = a
        for cell in row: set_cell_margins(cell)

    doc.add_page_break()
    
    # 3. CORE CAPABILITIES (100+ POINTS)
    h2 = doc.add_heading('2. CAPABILITIES MATRIX', level=1)
    for run in h2.runs: 
        run.font.name = 'Segoe UI'
        run.font.color.rgb = COLOR_PRIMARY
    
    caps_table = doc.add_table(rows=1, cols=2)
    caps_table.style = 'Table Grid'
    hdr = caps_table.rows[0].cells
    hdr[0].text = 'Domain'; hdr[1].text = 'Specific Skills'
    for cell in hdr:
        set_cell_background(cell, '1E293B')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    capabilities = [
        ['Local OS Control', 'Turbo Boost, Bloat Cleanup, Resource Monitoring.'],
        ['Core Engineering', 'Git-Nexus, Docker Sync, Database Migrations.'],
        ['AI Orchestration', 'Jules Background Agents, Automated Task Assembly.'],
        ['Visual Experience', 'Next-Gen Document & PDF Brand Generation.']
    ]
    for domain, skill in capabilities:
        row = caps_table.add_row().cells
        row[0].text = domain; row[1].text = skill
        for cell in row: set_cell_margins(cell)

    doc.add_paragraph("\n")
    h3 = doc.add_heading('3. VISION: ZERO FRICTION DEVELOPMENT', level=1)
    for run in h3.runs: run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph(
        'Integrating Python flexibility with Rust performance. '
        'Everything is a Skill. Everything is Automatable.'
    )

    # Save
    path = r'e:\Downloads\--ANTIGRAVITY store\IDE-optimus\PROJECT\outputs\NEXUS_Presentation_RU.docx'
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"PREMIUM_PRESENTATION_GENERATED: {path}")

if __name__ == "__main__":
    create_presentation()
