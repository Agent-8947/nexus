import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# DESIGN SYSTEM: MONOLITH LUXURY
COLOR_BG = RGBColor(0x05, 0x05, 0x05)
COLOR_TEXT_DIM = RGBColor(0x64, 0x74, 0x8B)
COLOR_TEXT_MAIN = RGBColor(0x0F, 0x17, 0x2A)
COLOR_ACCENT = RGBColor(0x00, 0x00, 0x00) # Pure Black for high contrast
PATH_COVER = r'C:\Users\MAC\.gemini\antigravity\brain\858d0bd1-44f8-418a-847c-07d05808974d\nexus_monolith_concept_1772197577106.png'

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_spacer(doc, points=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)

def create_premium_report():
    doc = Document()
    
    # 1. COVER PAGE
    if os.path.exists(PATH_COVER):
        doc.add_picture(PATH_COVER, width=Inches(6.0))
    
    add_spacer(doc, 40)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run('THE MONOLITH ARCHITECTURE')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.name = 'Segoe UI Semibold'
    run.font.color.rgb = COLOR_ACCENT
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('REDEFINING HIGH-PERFORMANCE MESSAGING // 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT_DIM
    run.font.name = 'Segoe UI'
    
    doc.add_page_break()

    # SECTION 1: EXECUTIVE SUMMARY
    h1 = doc.add_heading('EXECUTIVE SUMMARY', level=1)
    for run in h1.runs: 
        run.font.name = 'Segoe UI Semibold'
        run.font.color.rgb = COLOR_ACCENT
    
    p = doc.add_paragraph(
        "Project Nexus represents a paradigm shift in local machine infrastructure. "
        "By leveraging the memory safety of Rust and the rapid development cycle of Python, "
        "we have engineered a kernel capable of processing 100k messages per second on "
        "limited hardware (6GB RAM)."
    )
    p.style.font.name = 'Segoe UI'
    p.style.font.size = Pt(11)
    
    add_spacer(doc, 20)

    # TECHNOLOGY MATRIX TABLE
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    
    # Header styling
    hdr = table.rows[0].cells
    cols = ['LAYER', 'TECHNOLOGY', 'RATIONALE']
    for i, c in enumerate(cols):
        hdr[i].text = c
        run = hdr[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEXT_DIM
        set_cell_border(hdr[i], bottom={"sz": 4, "color": "64748B"})

    data = [
        ['Infrastructure', 'Rust (Tokio, Sled)', 'Sub-ms latency, No-GC spikes.'],
        ['Business Logic', 'Python 3.13 (FastHTML)', 'Rapid AI-integrated iteration.'],
        ['Orchestration', 'Go (PocketBase)', 'Lightweight discovery and auth.'],
        ['UI Layer', 'HTMX / Vanilla CSS', 'Zero-overhead dynamic interfaces.']
    ]
    
    for layer, tech, rat in data:
        row = table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = rat
        for cell in row:
            for p in cell.paragraphs:
                p.style.font.name = 'Segoe UI'
                p.style.font.size = Pt(10)
            set_cell_border(cell, bottom={"sz": 2, "color": "E2E8F0"})

    add_spacer(doc, 40)

    # SECTION 2: CRITICAL DECISIONS (ADR)
    h2 = doc.add_heading('CRITICAL DECISION LOG', level=1)
    for run in h2.runs: 
        run.font.name = 'Segoe UI Semibold'
        run.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph(
        "ADR-001: The selection of Rust over Java/C#. Rationale: Zero-cost abstractions "
        "and lack of a garbage collector are mandatory for the 100k throughput target "
        "on 6GB RAM machines where GC pauses would cause critical failure."
    , style='List Bullet')

    doc.add_paragraph(
        "ADR-002: Python logic integration. Rationale: While the core is Rust, "
        "top-level decisions require high agility and easy integration with modern AI models."
    , style='List Bullet')

    # Save
    output_path = r'e:\Downloads\--ANTIGRAVITY store\IDE-optimus\PROJECT\outputs\Architecture_Report_Nexus.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"ULTRA_PREMIUM_REPORT_GENERATED: {output_path}")

if __name__ == "__main__":
    create_premium_report()
