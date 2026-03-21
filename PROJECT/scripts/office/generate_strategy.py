import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# DESIGN SYSTEM: NEXUS MONOLITH (V3.0)
COLOR_BG = RGBColor(0x05, 0x05, 0x05)
COLOR_TEXT_DIM = RGBColor(0x64, 0x74, 0x8B)
COLOR_TEXT_MAIN = RGBColor(0x0F, 0x17, 0x2A)
COLOR_ACCENT = RGBColor(0x25, 0x63, 0xEB)
COLOR_GOLD = RGBColor(0xD4, 0xAF, 0x37)
PATH_COVER = r'C:\Users\MAC\.gemini\antigravity\brain\858d0bd1-44f8-418a-847c-07d05808974d\nexus_monolith_concept_1772197577106.png'

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level, color=COLOR_TEXT_MAIN):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Segoe UI Semibold'
        run.font.color.rgb = color
    return h

def create_strategy_report():
    doc = Document()
    
    # 1. COVER PAGE
    if os.path.exists(PATH_COVER):
        doc.add_picture(PATH_COVER, width=Inches(6.0))
    
    doc.add_paragraph("\n" * 2)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('STRATEGIC ROADMAP: NEXUS ASCENSION')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Segoe UI Semibold'
    run.font.color.rgb = COLOR_TEXT_MAIN
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('QUARTER 2 // 3-MONTH EXECUTION STRATEGY')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = COLOR_TEXT_DIM
    
    doc.add_page_break()

    # SECTION 1: VISION
    add_styled_heading(doc, '1. THE VISION: FROM LITE TO MONOLITH', level=1)
    doc.add_paragraph(
        "The next three months will focus on transitioning from a functional prototype to "
        "a market-ready, autonomous messaging ecosystem. Our strategy balances technical "
        "superiority (Rust) with operational agility (Python AI Agents)."
    )

    # 2. ROADMAP TABLE
    add_styled_heading(doc, '2. EXECUTION ROADMAP', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['PHASE', 'OBJECTIVES', 'DELIVERABLES']):
        hdr[i].text = txt
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        set_cell_border(hdr[i], bottom={"sz": 12, "color": "2563EB"})

    roadmap = [
        ['MONTH 1:\nCORE STABILITY', 
         "• Finalize Rust connectivity.\n• Named Pipes stabilization.\n• Basic Python IPC bridge.",
         "Working Core MVP (10k msg/s)."],
        ['MONTH 2:\nINTELLIGENCE',
         "• Bayesian metrics integration.\n• SKILL library expansion (800+).\n• Automated stress testing.",
         "Self-Optimizing System Status."],
        ['MONTH 3:\nECOSYSTEM',
         "• Omni-Dashboard finalization.\n• Telegram/Vercel automation.\n• Enterprise Documentation.",
         "Production-Ready Monolith."]
    ]
    
    for month, obj, dlv in roadmap:
        row = table.add_row().cells
        row[0].text = month
        row[1].text = obj
        row[2].text = dlv
        for cell in row:
            for p in cell.paragraphs:
                p.style.font.name = 'Segoe UI'
                p.style.font.size = Pt(10)

    doc.add_page_break()

    # SECTION 3: CORE PILLARS
    add_styled_heading(doc, '3. STRATEGIC PILLARS', level=1)
    
    pillars = [
        ("Hardware Agnosticism", "Optimization for 6GB RAM ensures scalability to any edge device."),
        ("Agentic Sovereignty", "Transitioning from manual triggers to autonomous agent decision-making."),
        ("Luxury DX", "Elite Developer Experience via premium dashboards and automated reporting.")
    ]
    
    for title, desc in pillars:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    # SECTION 4: RISK MITIGATION
    add_styled_heading(doc, '4. RISK MANAGEMENT', level=1, color=RGBColor(0xF4, 0x3F, 0x5E))
    
    risk_table = doc.add_table(rows=1, cols=2)
    hdr = risk_table.rows[0].cells
    hdr[0].text = 'RISK'; hdr[1].text = 'MITIGATION'
    
    risks = [
        ['Rust Learning Curve', 'Encapsulation of Rust code behind stable Python APIs.'],
        ['Windows Constraints', 'Aggressive use of nexus-system-control (Turbo Boost).'],
        ['Scope Creep', 'Strict adherence to Phase-gate deliverables defined in this doc.']
    ]
    for r, m in risks:
        row = risk_table.add_row().cells
        row[0].text = r; row[1].text = m

    # Save
    output_path = r'e:\Downloads\--ANTIGRAVITY store\IDE-optimus\PROJECT\outputs\Project_Strategy_3Months.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"STRATEGY_GENERATED: {output_path}")

if __name__ == "__main__":
    create_strategy_report()
