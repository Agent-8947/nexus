import os
from pathlib import Path
from docx import Document

def merge_docx_files(files, output_path):
    merged_doc = Document()
    for i, file_path in enumerate(files):
        print(f"  [-] Merging: {file_path}")
        sub_doc = Document(file_path)
        
        # Add heading if it's not the first file
        if i > 0:
            merged_doc.add_page_break()
            
        # Add filename as a title-like text for separation
        merged_doc.add_paragraph(f"--- Document: {os.path.basename(file_path)} ---", style='Heading 1')
        
        # Append elements
        for element in sub_doc.element.body:
            merged_doc.element.body.append(element)
            
    merged_doc.save(output_path)
    print(f"✅ Created: {output_path}")

def main():
    inbox_dir = Path("INBOX")
    archive_dir = Path("ARCHIVE/INBOX_backup")
    archive_dir.mkdir(parents=True, exist_ok=True)
    
    languages = ['EN', 'RU', 'UA']
    
    for lang in languages:
        print(f"Processing language: {lang}")
        lang_files = sorted([str(f) for f in inbox_dir.glob(f"NEXUS_{lang}_*.docx")])
        
        if not lang_files:
            print(f"  [!] No files found for {lang}")
            continue
            
        output_name = f"NEXUS_{lang}_COMBINED.docx"
        output_path = inbox_dir / output_name
        
        # Create a folder for the language if it doesn't exist
        lang_folder = inbox_dir / lang
        lang_folder.mkdir(exist_ok=True)
        
        # Merge
        merge_docx_files(lang_files, str(lang_folder / output_name))
        
        # Move original files to archive
        for f in lang_files:
            target = archive_dir / Path(f).name
            if target.exists():
                target.unlink() # Delete existing backup if needed or rename
            os.rename(f, str(target))
            print(f"  [>] Archived: {f}")

if __name__ == "__main__":
    main()
